"""Generate the four submission docs for PixLang.

Each doc has:
  - Title (Heading 1)
  - Subheadings (Heading 2)
  - Body paragraphs in 11pt Calibri, all black, justified
  - Code/grammar/pipeline blocks in 10pt Consolas with subtle background
  - Bulleted lists where appropriate
  - No em-dashes
"""
import os
import re

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "docs")
os.makedirs(OUT, exist_ok=True)

BLACK = RGBColor(0, 0, 0)


# ── Low-level helpers ────────────────────────────────────────────────────────

def _shade(paragraph, hex_fill):
    """Add a background fill to a paragraph (for code blocks)."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    pPr.append(shd)


def add_para(doc, text, bold=False, italic=False, size=11, font="Calibri",
             align=None, space_after=6, space_before=0, indent_left=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after  = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    if indent_left:
        pf.left_indent = Cm(indent_left)
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    run.bold = bold
    run.italic = italic
    return p


def add_heading(doc, text, level=1):
    sizes = {1: 18, 2: 13, 3: 11}
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(14 if level == 1 else 10)
    pf.space_after  = Pt(6)
    pf.keep_with_next = True
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(sizes[level])
    run.bold = True
    run.font.color.rgb = BLACK
    return p


def add_code_block(doc, lines, shaded=True):
    """One paragraph per line, in monospaced font with optional background."""
    for line in lines:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_after  = Pt(0)
        pf.space_before = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.1
        pf.left_indent = Cm(0.5)
        if shaded:
            _shade(p, "F2F2F2")
        run = p.add_run(line if line else " ")  # nbsp keeps blank shaded
        run.font.name = "Consolas"
        run.font.size = Pt(10)
        run.font.color.rgb = BLACK
    # trailing whitespace after the block
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


REPO_URL = "https://github.com/MuhammadMehroz786/Pixel-Lang"


def add_repo_link(doc, label="Project Repository"):
    """Insert 'Project Repository: <clickable url>' below the title."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(10)

    lbl_run = p.add_run(f"{label}: ")
    lbl_run.font.name = "Calibri"
    lbl_run.font.size = Pt(11)
    lbl_run.font.color.rgb = BLACK
    lbl_run.bold = True

    # Hyperlink
    part = doc.part
    r_id = part.relate_to(
        REPO_URL,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")  # 11pt = 22 half-points
    rPr.append(sz)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    new_run.append(rPr)
    text_el = OxmlElement("w:t")
    text_el.text = REPO_URL
    new_run.append(text_el)
    hyperlink.append(new_run)
    p._p.append(hyperlink)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        pf = p.paragraph_format
        pf.space_after = Pt(2)
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = BLACK


def justify(doc, text):
    """Add a justified body paragraph; collapse single newlines into spaces."""
    text = re.sub(r"\s+", " ", text).strip()
    return add_para(doc, text, align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                    space_after=8)


def set_margins(doc, cm=2.5):
    for section in doc.sections:
        section.top_margin    = Cm(cm)
        section.bottom_margin = Cm(cm)
        section.left_margin   = Cm(cm)
        section.right_margin  = Cm(cm)


def new_doc():
    doc = Document()
    set_margins(doc)
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = BLACK
    return doc


# ── Document 1: Team Reflection ──────────────────────────────────────────────

def make_team_reflection():
    doc = new_doc()
    add_heading(doc, "Team Reflection", level=1)
    add_repo_link(doc)

    justify(doc, """
        Working on PixLang gave us a much clearer picture of what actually
        happens between writing a piece of source code and seeing a result
        on screen. Splitting the compiler into six discrete phases (lexer,
        parser, semantic analyser, IR generator, optimiser, and code
        generator) forced us to think about each stage in isolation rather
        than as one big translation step.
    """)

    add_heading(doc, "Team Contributions", level=2)

    add_para(doc, "Mehroz Muneer (23K-0748)", bold=True, space_after=2)
    justify(doc, """
        Led the front-end of the compiler. Built the hand-written DFA
        lexer and the recursive-descent parser, designed the AST node
        hierarchy in ast_nodes.py, and produced the EBNF grammar that the
        parser implements. Set up the CLI driver in compiler.py, integrated
        the six phases into a single pipeline, and authored the --debug
        output that prints every intermediate representation. Wrote and
        debugged most of the test programs in tests/ as the language took
        shape.
    """)

    add_para(doc, "Dua Shafiq (23K-0825)", bold=True, space_after=2)
    justify(doc, """
        Owned the middle and back-end of the compiler. Wrote the semantic
        analyser including the scoped symbol table and the type system
        (INT, FLOAT, COLOR, STRING), handled CANVAS-before-drawing
        validation and palette resolution, and produced the line-accurate
        semantic error messages. Implemented the IR generator that lowers
        the AST to Three-Address Code and the two optimiser passes
        (constant folding and dead-variable elimination) that make the
        --debug instruction counts drop visibly.
    """)

    add_para(doc, "Kashan Hamid (23K-0747)", bold=True, space_after=2)
    justify(doc, """
        Owned code generation, tooling, and the user-facing parts of the
        project. Built the TAC interpreter / virtual machine in codegen.py
        that drives the Pillow image library, the interactive REPL, and
        the Streamlit playground in ui/app.py. Generated the parse-tree,
        DFA, and symbol-table diagrams under diagrams/, organised the
        project into its final src / ui / tests / outputs / docs layout,
        and wrote and maintained the README.
    """)

    add_heading(doc, "Challenges", level=2)
    justify(doc, """
        The hardest parts were the parser and the semantic analyser.
        Recursive descent looks straightforward in theory, but handling
        operator precedence, palette indexing, and statements that span
        multiple lines (LOOP and IF blocks) required careful planning.
        Semantic analysis was tricky in a different way: keeping the
        symbol table consistent across nested scopes and producing useful
        error messages took several rewrites before it felt right.
    """)

    add_heading(doc, "What We Learned and Future Work", level=2)
    justify(doc, """
        Generating Three-Address Code and watching the optimiser fold
        constants and remove dead variables was the most satisfying part.
        It made the theory from class concrete: we could literally see the
        instruction count drop in the --debug output. If we extended
        PixLang we would add user-defined functions, nested palettes, and
        a richer optimiser (common subexpression elimination, loop-invariant
        code motion). We would also like to expand the Streamlit playground
        into a proper IDE with syntax highlighting and a live preview.
    """)

    doc.save(os.path.join(OUT, "Team_Reflection.docx"))


# ── Document 2: Compiler Architecture ────────────────────────────────────────

def make_compiler_architecture():
    doc = new_doc()
    add_heading(doc, "Compiler Architecture", level=1)
    add_repo_link(doc)

    justify(doc, """
        PixLang compiles a .pix source file into a PNG image through a
        six-phase pipeline. Each phase is implemented in its own Python
        module under src/, and compiler.py wires them together.
    """)

    add_heading(doc, "Pipeline", level=2)
    add_code_block(doc, [
        "Source Code",
        "    -> Lexer            (Phase 1)",
        "    -> Parser           (Phase 2)",
        "    -> Semantic Analyser (Phase 3)",
        "    -> IR Generator     (Phase 4)",
        "    -> Optimiser        (Phase 5)",
        "    -> Code Generator   (Phase 6)",
        "    -> output.png",
    ])

    add_heading(doc, "Phase 1: Lexer", level=2)
    justify(doc, """
        A hand-written deterministic finite automaton converts the input
        character stream into a list of tokens (keywords, identifiers,
        integer and float literals, color literals, strings, operators,
        and punctuation). Comments beginning with ';' and whitespace are
        discarded.
    """)

    add_heading(doc, "Phase 2: Parser", level=2)
    justify(doc, """
        A recursive-descent parser consumes the token stream and builds
        an Abstract Syntax Tree using the node classes defined in
        ast_nodes.py. The parser enforces the grammar of statements,
        expressions, and control-flow blocks, and reports the line and
        column of any syntax error.
    """)

    add_heading(doc, "Phase 3: Semantic Analyser", level=2)
    justify(doc, """
        Walks the AST, builds a scoped symbol table, infers expression
        types (INT, FLOAT, COLOR, STRING), validates that drawing commands
        appear after a CANVAS declaration, and checks palette references.
        Each lexical scope corresponds to a LOOP or IF block.
    """)

    add_heading(doc, "Phase 4: IR Generator", level=2)
    justify(doc, """
        Lowers the validated AST into Three-Address Code, a flat linear
        sequence of simple instructions that resemble assembly. Control
        flow is encoded with explicit labels and conditional jumps.
    """)

    add_heading(doc, "Phase 5: Optimiser", level=2)
    justify(doc, """
        Performs two passes on the TAC: constant folding, which evaluates
        constant sub-expressions at compile time, and dead variable
        elimination, which removes temporaries that are written but never
        read. The number of instructions before and after optimisation is
        reported with the --debug flag.
    """)

    add_heading(doc, "Phase 6: Code Generator", level=2)
    justify(doc, """
        Interprets the optimised TAC instructions on a small virtual
        machine that drives the Pillow image library. Drawing instructions
        place rectangles, circles, lines, and individual pixels onto the
        canvas; the SAVE instruction writes the result to a PNG file.
    """)

    doc.save(os.path.join(OUT, "CompilerArchitecture.docx"))


# ── Document 3: Language Reference Manual ────────────────────────────────────

def make_language_reference():
    doc = new_doc()
    add_heading(doc, "Language Reference Manual", level=1)
    add_repo_link(doc)

    justify(doc, """
        PixLang is a small domain-specific language for procedurally
        generating pixel art. A program describes a canvas, an optional
        set of named palettes, and a sequence of drawing commands. Loops,
        conditionals, variables, and arithmetic expressions allow patterns
        to be described compactly rather than pixel by pixel.
    """)

    add_heading(doc, "Keywords", level=2)
    add_code_block(doc, [
        "CANVAS    PALETTE   FILL      VAR       LOOP",
        "WHILE     TIMES     ENDLOOP   IF        ELSE",
        "ENDIF     RECT      CIRCLE    LINE      PIXEL",
        "SAVE      RADIUS    COLOR     MOD",
    ])

    add_heading(doc, "Lexical Tokens", level=2)
    add_bullets(doc, [
        "IDENTIFIER: matches [A-Za-z_][A-Za-z0-9_]*",
        "INT and FLOAT: standard decimal literals",
        "COLOR_LIT: hexadecimal in the form #RGB or #RRGGBB",
        "STRING: double-quoted text, used with SAVE",
        "COMMENT: starts with ';' and runs to the end of the line",
    ])

    add_heading(doc, "Drawing Commands", level=2)
    add_code_block(doc, [
        "FILL    color",
        "RECT    x, y, width, height, color",
        "CIRCLE  cx cy RADIUS r COLOR color",
        "LINE    x1 y1 x2 y2 COLOR color",
        "PIXEL   x  y  COLOR color",
    ])

    add_heading(doc, "Variables and Arithmetic", level=2)
    justify(doc, """
        Variables are declared with VAR and reassigned without VAR.
        Arithmetic supports +, -, *, /, MOD with full operator precedence
        and parentheses. Comparisons (<, >, <=, >=, ==, !=) yield numeric
        truth values used inside conditions.
    """)

    add_heading(doc, "Control Flow", level=2)
    add_code_block(doc, [
        "LOOP WHILE condition:",
        "    statements",
        "ENDLOOP",
        "",
        "LOOP TIMES n:",
        "    statements",
        "ENDLOOP",
        "",
        "IF condition:",
        "    then_statements",
        "ELSE:",
        "    else_statements",
        "ENDIF",
    ])

    add_heading(doc, "Palettes", level=2)
    justify(doc, """
        Palettes are declared with PALETTE name: #hex, #hex, ... and
        indexed as name[i] where i is any integer expression. Palette
        indices are zero-based.
    """)

    add_heading(doc, "Grammar (Simplified)", level=2)
    add_code_block(doc, [
        "program     ::= statement*",
        "statement   ::= canvas_stmt | palette_stmt | draw_stmt",
        "              | var_decl | assign | loop_stmt",
        "              | if_stmt | save_stmt",
        "loop_stmt   ::= 'LOOP' ('WHILE' expr | 'TIMES' expr) ':'",
        "                statement* 'ENDLOOP'",
        "if_stmt     ::= 'IF' expr ':' statement*",
        "                ['ELSE' ':' statement*] 'ENDIF'",
        "expr        ::= comparison",
        "comparison  ::= addition (relop addition)?",
        "addition    ::= term (('+' | '-') term)*",
        "term        ::= factor (('*' | '/' | 'MOD') factor)*",
    ])

    add_heading(doc, "Output and Optimisation", level=2)
    justify(doc, """
        Programs end with SAVE "file.png" to write the rendered canvas.
        The compiler performs constant folding and dead variable
        elimination on the intermediate representation before code
        generation. With the --debug flag the compiler prints every
        intermediate phase and the before/after instruction counts of
        the optimiser.
    """)

    doc.save(os.path.join(OUT, "LanguageReferenceManual.docx"))


# ── Document 4: Test Suite ───────────────────────────────────────────────────

def _test_entry(doc, n, name, canvas, body):
    add_heading(doc, f"Test {n}: {name}  ({canvas})", level=2)
    justify(doc, body)


def make_test_suite():
    doc = new_doc()
    add_heading(doc, "Test Suite", level=1)
    add_repo_link(doc)

    justify(doc, """
        PixLang ships with six test programs in tests/, each chosen to
        exercise a different combination of language features. Running a
        test produces a PNG in outputs/, which can be compared against the
        reference image bundled in the same folder. All six programs are
        expected to compile and produce identical pixel output across runs.
    """)

    _test_entry(doc, 1, "test1_concentric.pix", "64 x 64", """
        Exercises palettes, palette indexing, VAR declarations, arithmetic
        with parentheses and division, the MOD operator, and a LOOP WHILE
        block. Draws concentric rectangles around a central circle.
        Expected output: nested coloured squares on a dark background.
    """)

    _test_entry(doc, 2, "test2_flag.pix", "90 x 60", """
        The simplest program in the suite. No variables, no loops: just
        FILL plus three RECT commands and a SAVE. Used as a sanity check
        that the parser and codegen handle linear straight-line programs
        correctly. Expected output: three horizontal stripes (red, white,
        green).
    """)

    _test_entry(doc, 3, "test3_checkerboard.pix", "64 x 64", """
        Exercises nested LOOP TIMES blocks together with an IF / ELSE
        inside the inner loop. Demonstrates that scope handling in the
        semantic analyser is correct across multiple nested blocks.
        Expected output: a regular checkerboard pattern.
    """)

    _test_entry(doc, 4, "test4_starburst.pix", "128 x 128", """
        Exercises the LINE primitive together with palette indexing.
        Eight LINE statements draw rays from the centre of the canvas to
        each of the eight compass directions, each tinted with a different
        colour from an 8-entry palette, with a small CIRCLE finishing the
        centre. Expected output: a multicoloured starburst on a black
        background.
    """)

    add_heading(doc, "Test 5: test5_error_demo.pix", level=2)
    justify(doc, """
        A two-mode test. By default it compiles successfully (a fallback
        program that draws three coloured rectangles) so that the run-all
        script does not fail. The file also contains five commented-out
        blocks, each demonstrating a different error class. Uncomment one
        block at a time to see the corresponding diagnostic.
    """)
    add_bullets(doc, [
        "Lexer error: invalid character ('@')",
        "Parser error: missing ENDLOOP",
        "Semantic error: drawing command before CANVAS",
        "Semantic error: undeclared variable",
        "Semantic error: wrong SAVE extension (.bmp)",
    ])

    _test_entry(doc, 6, "test6_character.pix", "256 x 320", """
        A blocky, Roblox-style figure built entirely from RECT and PIXEL
        commands. Demonstrates that the language is expressive enough for
        hand-authored sprite art at a larger canvas size, and acts as a
        visually striking demo during the viva. Expected output: a
        yellow-headed, blue-shirted, red-trousered character on a dark
        background.
    """)

    add_heading(doc, "Running the Full Suite", level=2)
    add_code_block(doc, [
        "python compiler.py tests/test1_concentric.pix",
        "python compiler.py tests/test2_flag.pix",
        "python compiler.py tests/test3_checkerboard.pix",
        "python compiler.py tests/test4_starburst.pix",
        "python compiler.py tests/test5_error_demo.pix",
        "python compiler.py tests/test6_character.pix",
    ])

    justify(doc, """
        Each command writes its image into outputs/ and prints
        "Compilation successful." on success. Add --debug to any command
        to see the token stream, AST, symbol table, raw and optimised
        TAC, and the codegen log for that program.
    """)

    doc.save(os.path.join(OUT, "TestSuite.docx"))


# ── Driver ───────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    make_team_reflection()
    print(f"  wrote {OUT}/Team_Reflection.docx")
    make_compiler_architecture()
    print(f"  wrote {OUT}/CompilerArchitecture.docx")
    make_language_reference()
    print(f"  wrote {OUT}/LanguageReferenceManual.docx")
    make_test_suite()
    print(f"  wrote {OUT}/TestSuite.docx")
